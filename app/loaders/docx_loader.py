from docx import Document
from langchain_core.documents import Document as LCDocument

from app.loaders.base_loader import BaseLoader

class DOCXLoader(BaseLoader):

    def extract_text(self, file_path):
        doc = Document(file_path)

        text = "\n".join(para.text for para in doc.paragraphs)

        return [LCDocument(
            page_content=text,
            metadata={"source": file_path}
        )]